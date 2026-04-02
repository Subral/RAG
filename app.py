from __future__ import annotations
import os
import re
import sqlite3
import threading
from pathlib import Path
from typing import Iterable, List, Tuple, Dict
from contextlib import asynccontextmanager
import subprocess
import sys
import tempfile
import textwrap
import json
import faiss
import numpy as np
import ollama
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import pandas as pd
import uuid
from fastapi import UploadFile, File, Form
from functools import lru_cache

DEBUG = os.getenv("DEBUG", "false").lower() == "true"

def dlog(*args):
    if DEBUG:
        print(*args, flush=True)

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class QueryResponse(BaseModel):
    answer: str


EMBEDDING_MODEL = "qwen3-embedding:4b"

LANGUAGE_MODEL = "qwen2.5-coder:7b"
CODE_MODEL = "qwen2.5:14b"   
NARRATE_MODEL = "qwen2.5:14b"

MAX_CHARS_PER_CHUNK = 300
EMBED_BATCH_SIZE = 16
TABULAR_EXTENSIONS = {".csv", ".xlsx", ".xls"}


def is_tabular(path: str) -> bool:
    ext = Path(path).suffix.lower()
    result = ext in TABULAR_EXTENSIONS
    print(f"[DEBUG][is_tabular] path='{path}' | ext='{ext}' | is_tabular={result}")
    return result

def build_schema_context(df: pd.DataFrame) -> str:
    print(f"[DEBUG][build_schema_context] Building schema for df with shape={df.shape}")
    lines = [f"DataFrame shape: {df.shape[0]} rows x {df.shape[1]} columns\n"]
    lines.append("Columns:")
    for col in df.columns:
        dtype = str(df[col].dtype)
        nulls = int(df[col].isna().sum())
        is_numeric = pd.api.types.is_numeric_dtype(df[col])
        print(f"[DEBUG][build_schema_context]   col='{col}' | dtype={dtype} | is_numeric={is_numeric}")
        if is_numeric:
            sample = (
                f"min={df[col].min()}, max={df[col].max()}, "
                f"mean={df[col].mean():.2f}"
            )
        else:
            unique_vals = df[col].dropna().unique()[:6].tolist()
            sample = f"sample unique values: {unique_vals}"
        col_line = f"  - '{col}' | dtype: {dtype} | nulls: {nulls} | {sample}"
        print(f"[DEBUG][build_schema_context]   {col_line}")
        lines.append(col_line)
    schema = "\n".join(lines)
    print(f"[DEBUG][build_schema_context] Schema built successfully ({len(lines)-2} columns described)")
    return schema


def load_dataframe(path: str) -> pd.DataFrame:
    ext = Path(path).suffix.lower()
    print(f"[DEBUG][load_dataframe] Loading '{path}' | ext='{ext}'")
    if ext == ".csv":
        df = pd.read_csv(path)
    elif ext in (".xlsx", ".xls"):
        df = pd.read_excel(path)
    else:
        raise ValueError(f"Unsupported tabular format: {ext}")
    print(f"[DEBUG][load_dataframe] Loaded successfully | shape={df.shape} | columns={df.columns.tolist()}")
    return df



SANDBOX_TIMEOUT = 10  

_SANDBOX_WRAPPER = """
import pandas as pd
import json, sys

df = pd.read_csv({csv_path!r})

# ---- LLM-generated code below ----
{user_code}
# ---- end of LLM code ----

try:
    if isinstance(result, pd.DataFrame):
        out = result.to_dict(orient="records")
    elif isinstance(result, pd.Series):
        out = result.to_dict()
    elif isinstance(result, tuple):          
        out = list(result)
    elif hasattr(result, 'item'):            # numpy scalar
        out = result.item()
    else:
        out = result
    print(json.dumps({{"ok": True, "result": out}}))
except NameError:
    print(json.dumps({{"ok": False, "error": "LLM code did not assign to `result`"}}))
"""

def run_pandas_code_sandbox(df: pd.DataFrame, code: str) -> Tuple[bool, str]:
    print(f"[DEBUG][run_pandas_code_sandbox] Starting sandbox execution")
    print(f"[DEBUG][run_pandas_code_sandbox] Code to run:\n{'='*50}\n{code}\n{'='*50}")

    with tempfile.TemporaryDirectory() as tmp:
        csv_path = os.path.join(tmp, "data.csv")
        script_path = os.path.join(tmp, "script.py")

        df.to_csv(csv_path, index=False)
        print(f"[DEBUG][run_pandas_code_sandbox] df serialized to temp CSV: '{csv_path}'")

        script = _SANDBOX_WRAPPER.format(
            csv_path=csv_path,
            user_code=textwrap.indent(code, ""),
        )
        with open(script_path, "w") as f:
            f.write(script)
        print(f"[DEBUG][run_pandas_code_sandbox] Script written to: '{script_path}'")

        print(f"[DEBUG][run_pandas_code_sandbox] Launching subprocess | python={sys.executable} | timeout={SANDBOX_TIMEOUT}s")
        try:
            proc = subprocess.run(
                [sys.executable, script_path], 
                capture_output=True,
                text=True,
                timeout=SANDBOX_TIMEOUT,
            )
        except subprocess.TimeoutExpired:
            print(f"[DEBUG][run_pandas_code_sandbox] FAILED: subprocess timed out after {SANDBOX_TIMEOUT}s")
            return False, "Execution timed out after 10 seconds."

        stdout = proc.stdout.strip()
        stderr = proc.stderr.strip()
        print(f"[DEBUG][run_pandas_code_sandbox] Subprocess finished | returncode={proc.returncode}")
        print(f"[DEBUG][run_pandas_code_sandbox] stdout: {stdout[:300] if stdout else '(empty)'}")
        if stderr:
            print(f"[DEBUG][run_pandas_code_sandbox] stderr: {stderr[:500]}")

        if not stdout:
            print(f"[DEBUG][run_pandas_code_sandbox] FAILED: no stdout output")
            return False, f"No output from script. stderr: {stderr}"

        try:
            payload = json.loads(stdout)
        except json.JSONDecodeError as e:
            print(f"[DEBUG][run_pandas_code_sandbox] FAILED: could not parse stdout as JSON: {e}")
            return False, f"Could not parse output: {stdout}"

        if payload.get("ok"):
            result_preview = str(payload["result"])[:300]
            print(f"[DEBUG][run_pandas_code_sandbox] SUCCESS | result preview: {result_preview}")
            return True, json.dumps(payload["result"])
        else:
            err = payload.get("error", stderr)
            print(f"[DEBUG][run_pandas_code_sandbox] FAILED: script reported error: {err}")
            return False, err


def generate_pandas_code(
    query: str,
    schema_context: str,
    previous_error: str | None = None,
) -> str:
    print(f"[DEBUG][generate_pandas_code] Generating pandas code for query: '{query}'")
    if previous_error:
        print(f"[DEBUG][generate_pandas_code] Previous error being fed back to LLM: {previous_error[:200]}")

    error_block = ""
    if previous_error:
        error_block = f"\nYour previous attempt failed with this error:\n{previous_error}\nFix it.\n"

    prompt = f"""You are a pandas code generator.

CRITICAL: The dataframe columns are EXACTLY as listed in the schema below.
Do NOT guess or invent column names. Use only what is shown.

Schema:
{schema_context}

Rules:
1. df is already loaded. Do NOT reload it.
2. Assign final answer to `result`.
3. No imports, no print(), no file operations.
4. Code only — no explanation, no markdown fences.

{"Previous attempt failed: " + previous_error if previous_error else ""}

Question: {query}

Code:"""

    print(f"[DEBUG][generate_pandas_code] Calling LLM (model={CODE_MODEL})...")
    response = ollama.chat(
        model=CODE_MODEL,
        messages=[{"role": "user", "content": prompt}],
    )
    raw = response["message"]["content"].strip()
    print(f"[DEBUG][generate_pandas_code] Raw LLM response:\n{'-'*40}\n{raw}\n{'-'*40}")

    cleaned = re.sub(r"^```(?:python)?\n?", "", raw)
    cleaned = re.sub(r"\n?```$", "", cleaned).strip()
    if cleaned != raw:
        print(f"[DEBUG][generate_pandas_code] Stripped markdown fences from LLM output")

    # Strip any leading prose lines the LLM adds before actual code.
    PYTHON_START_KEYWORDS = (
        "result", "df", "import", "if", "for", "while", "return",
        "#", "try", "with", "def", "class", "raise", "print", "pd",
    )
    code_lines = []
    code_started = False
    for line in cleaned.split("\n"):
        stripped = line.strip()
        if not code_started:
            is_python = (
                any(stripped.startswith(kw) for kw in PYTHON_START_KEYWORDS)
                or bool(re.match(r"^[a-z_].*=", stripped))
                or stripped.startswith(("#", "(", "[", "{"))
            )
            if is_python:
                code_started = True
            else:
                print(f"[DEBUG][generate_pandas_code] Skipping prose line: '{stripped}'")
        if code_started:
            code_lines.append(line)
    cleaned = "\n".join(code_lines).strip()

    print(f"[DEBUG][generate_pandas_code] Final code to execute ({len(cleaned)} chars):\n{cleaned}")
    return cleaned


def query_tabular_file(
    df: pd.DataFrame,
    schema_context: str,
    query: str,
    max_retries: int = 3,
) -> str:
    print(f"[DEBUG][query_tabular_file] Starting tabular pipeline | query='{query}' | max_retries={max_retries}")
    last_error = None

    for attempt in range(1, max_retries + 1):
        print(f"\n[DEBUG][query_tabular_file] -- Attempt {attempt}/{max_retries} --")
        code = generate_pandas_code(query, schema_context, previous_error=last_error)
        success, result_or_error = run_pandas_code_sandbox(df, code)

        if success:
            print(f"[DEBUG][query_tabular_file] Attempt {attempt} SUCCEEDED, proceeding to narration")
            return narrate_result(query, result_or_error)
        else:
            print(f"[DEBUG][query_tabular_file] Attempt {attempt} FAILED: {result_or_error[:200]}")
            last_error = result_or_error

    print(f"[DEBUG][query_tabular_file] All {max_retries} attempts exhausted")
    return f"I was unable to compute an answer after {max_retries} attempts. Last error: {last_error}"


def narrate_result(query: str, result_json: str) -> str:
    print(f"[DEBUG][narrate_result] Narrating result for query: '{query}'")
    print(f"[DEBUG][narrate_result] Result JSON: {result_json[:300]}")

    prompt = f"""You are a data reporting tool. Report ONLY what the numbers say.

    STRICT RULES:
    1. Use ONLY the data result below. Nothing else.
    2. Do NOT add context, suggestions, or interpretation.
    3. Do NOT say things like "this suggests" or "this indicates".
    4. If the result is a single number, report just that number with its label.
    5. Do NOT mention data quality, sample size, or methodology.

    User question: {query}

    Data result (JSON):
    {result_json}

    Answer:"""

    print(f"[DEBUG][narrate_result] Calling LLM for narration (model={NARRATE_MODEL})...")
    response = ollama.chat(
        model=NARRATE_MODEL,
        messages=[{"role": "user", "content": prompt}],
    )
    answer = response["message"]["content"].strip()
    print(f"[DEBUG][narrate_result] Narration complete | answer preview: '{answer[:200]}'")
    return answer

def get_user_workspace(user_id: str) -> Path:
    base = Path("user_data") / user_id
    base.mkdir(parents=True, exist_ok=True)
    print(f"[DEBUG][get_user_workspace] Workspace for user '{user_id}': '{base}'")
    return base


def get_paths_for_pdf(pdf_path: str, user_id: str) -> tuple[str, str, str]:
    pdf_name = Path(pdf_path).stem
    folder = get_user_workspace(user_id) / pdf_name
    folder.mkdir(parents=True, exist_ok=True)
    db_path = folder / f"{pdf_name}_chunks.db"
    index_path = folder / f"{pdf_name}_hnsw.index"
    print(f"[DEBUG][get_paths_for_pdf] pdf_name='{pdf_name}' | db='{db_path}' | index='{index_path}'")
    return pdf_name, str(db_path), str(index_path)


def build_user_indexes(user_id: str, file_paths: List[str]):
    print(f"[DEBUG][build_user_indexes] Building indexes for {len(file_paths)} file(s): {file_paths}")
    all_indexes = []
    all_conns = []
    for file_path in file_paths:
        print(f"[DEBUG][build_user_indexes] Processing file: '{file_path}'")
        pdf_name, db_path, index_path = get_paths_for_pdf(file_path, user_id)
        index, conn = build_or_load_index(
            source_path=file_path,
            db_path=db_path,
            index_path=index_path,
        )
        all_indexes.append(index)
        all_conns.append(conn)
        print(f"[DEBUG][build_user_indexes] Index ready for '{pdf_name}' | total vectors={index.ntotal}")
    print(f"[DEBUG][build_user_indexes] All {len(all_indexes)} indexes ready")
    return all_indexes, all_conns


def _clean_text(s: str) -> str:
    s = s.replace("\u00ad", "")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\s+\n", "\n", s)
    return s.strip()


def load_pdf_text(path: str, password: str | None = None) -> List[str]:
    import fitz
    print(f"[DEBUG][load_pdf_text] Opening PDF: '{path}'")
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(path)
    doc = fitz.open(p.as_posix())
    print(f"[DEBUG][load_pdf_text] PDF opened | pages={len(doc)} | encrypted={doc.needs_pass}")
    if doc.needs_pass:
        if not password:
            raise ValueError("PDF is encrypted and no password was provided.")
        if not doc.authenticate(password):
            raise ValueError("Invalid PDF password.")
        print(f"[DEBUG][load_pdf_text] PDF authenticated successfully")
    lines: List[str] = []
    for i, page in enumerate(doc):
        text = page.get_text("text")
        text = _clean_text(text)
        if text:
            page_lines = [ln for ln in text.splitlines() if ln.strip()]
            lines.extend(page_lines)
            print(f"[DEBUG][load_pdf_text]   page {i+1}: {len(page_lines)} lines extracted")
    doc.close()
    print(f"[DEBUG][load_pdf_text] Total lines extracted from PDF: {len(lines)}")
    return lines


def load_docx_text(path: str) -> List[str]:
    import docx
    print(f"[DEBUG][load_docx_text] Opening DOCX: '{path}'")
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(path)
    doc = docx.Document(p.as_posix())
    lines: List[str] = []
    for para in doc.paragraphs:
        t = _clean_text(para.text)
        if t:
            lines.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                t = _clean_text(cell.text)
                if t:
                    lines.append(t)
    print(f"[DEBUG][load_docx_text] Total lines extracted from DOCX: {len(lines)}")
    return lines


def iter_lines(path: str) -> Iterable[str]:
    ext = Path(path).suffix.lower()
    print(f"[DEBUG][iter_lines] Iterating lines from '{path}' | ext='{ext}'")
    if ext == ".pdf":
        for ln in load_pdf_text(path):
            yield _clean_text(ln)
    elif ext == ".docx":
        for ln in load_docx_text(path):
            yield _clean_text(ln)
    elif ext in (".txt", ".md"):
        line_count = 0
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            for line in f:
                line = _clean_text(line.rstrip("\n"))
                if line:
                    line_count += 1
                    yield line
        print(f"[DEBUG][iter_lines] Total lines yielded from TXT/MD: {line_count}")
    else:
        raise ValueError(f"Unsupported file type for RAG: {ext}")


def iter_chunks_from_lines(
    line_iter: Iterable[str], max_chars: int = MAX_CHARS_PER_CHUNK
) -> Iterable[str]:
    buffer = ""
    chunk_count = 0
    for line in line_iter:
        if not line:
            continue
        if len(buffer) + len(line) + 1 <= max_chars:
            buffer += line + " "
        else:
            if buffer:
                chunk_count += 1
                yield buffer.strip()
            buffer = line + " "
    if buffer:
        chunk_count += 1
        yield buffer.strip()
    print(f"[DEBUG][iter_chunks_from_lines] Total chunks produced: {chunk_count}")


def count_chunks(path: str, max_chars: int = MAX_CHARS_PER_CHUNK) -> int:
    n = 0
    buffer = ""
    for line in iter_lines(path):
        if len(buffer) + len(line) + 1 <= max_chars:
            buffer += line + " "
        else:
            if buffer:
                n += 1
            buffer = line + " "
    if buffer:
        n += 1
    print(f"[DEBUG][count_chunks] Estimated chunk count for '{path}': {n}")
    return n


def init_db(db_path: str) -> sqlite3.Connection:
    print(f"[DEBUG][init_db] Initializing SQLite DB at '{db_path}'")
    conn = sqlite3.connect(db_path)
    cur = conn.cursor()
    cur.execute(
        "CREATE TABLE IF NOT EXISTS docs (id INTEGER PRIMARY KEY, text TEXT NOT NULL)"
    )
    conn.commit()
    print(f"[DEBUG][init_db] DB initialized successfully")
    return conn


def insert_chunks(conn: sqlite3.Connection, start_id: int, chunks: List[str]) -> None:
    cur = conn.cursor()
    rows = [(start_id + i, chunk) for i, chunk in enumerate(chunks)]
    cur.executemany("INSERT OR REPLACE INTO docs (id, text) VALUES (?, ?)", rows)
    conn.commit()
    print(f"[DEBUG][insert_chunks] Inserted {len(rows)} chunks | start_id={start_id} | end_id={start_id + len(rows) - 1}")


def fetch_chunk_by_id(conn: sqlite3.Connection, idx: int) -> str:
    cur = conn.cursor()
    cur.execute("SELECT text FROM docs WHERE id=?", (idx,))
    row = cur.fetchone()
    result = row[0] if row else ""
    if not result:
        print(f"[DEBUG][fetch_chunk_by_id] WARNING: no chunk found for id={idx}")
    return result


def embed_batch(text_batch: List[str]) -> np.ndarray:
    print(f"[DEBUG][embed_batch] Embedding {len(text_batch)} texts | model='{EMBEDDING_MODEL}'")
    try:
        resp = ollama.embed(model=EMBEDDING_MODEL, input=text_batch)
    except Exception as e:
        print(f"[DEBUG][embed_batch] WARNING: embedding failed ({e}), splitting batch and retrying...")
        if len(text_batch) == 1:
            raise e
        mid = len(text_batch) // 2
        left = embed_batch(text_batch[:mid])
        right = embed_batch(text_batch[mid:])
        return np.vstack([left, right])
    embs = np.array(resp["embeddings"], dtype="float32")
    faiss.normalize_L2(embs)
    print(f"[DEBUG][embed_batch] Embeddings produced | shape={embs.shape}")
    return embs


def embed_in_batches(
    chunk_iter: Iterable[str], batch_size: int = EMBED_BATCH_SIZE
) -> Iterable[Tuple[List[str], np.ndarray]]:
    print(f"[DEBUG][embed_in_batches] Starting batched embedding | batch_size={batch_size}")
    batch: List[str] = []
    total_embedded = 0
    for chunk in chunk_iter:
        batch.append(chunk)
        if len(batch) >= batch_size:
            embs = embed_batch(batch)
            total_embedded += len(batch)
            print(f"[DEBUG][embed_in_batches] Yielding batch | cumulative_embedded={total_embedded}")
            yield batch, embs
            batch = []
    if batch:
        embs = embed_batch(batch)
        total_embedded += len(batch)
        print(f"[DEBUG][embed_in_batches] Yielding final batch | cumulative_embedded={total_embedded}")
        yield batch, embs
    print(f"[DEBUG][embed_in_batches] Batched embedding complete | total_embedded={total_embedded}")


def embed_query(query: str) -> np.ndarray:
    print(f"[DEBUG][embed_query] Embedding query: '{query[:100]}'")
    resp = ollama.embed(model=EMBEDDING_MODEL, input=[query])
    q_emb = np.array(resp["embeddings"][0], dtype="float32").reshape(1, -1)
    faiss.normalize_L2(q_emb)
    print(f"[DEBUG][embed_query] Query embedding shape: {q_emb.shape}")
    return q_emb


def build_or_load_index(
    source_path: str,
    db_path: str,
    index_path: str,
) -> Tuple[faiss.Index, sqlite3.Connection]:
    print(f"[DEBUG][build_or_load_index] source='{source_path}'")
    
    if os.path.exists(index_path) and os.path.exists(db_path):
        print(f"[DEBUG][build_or_load_index] Existing index + DB found, loading from disk...")
        index = faiss.read_index(index_path)
        conn = sqlite3.connect(db_path)
        print(f"[DEBUG][build_or_load_index] Loaded index | total vectors={index.ntotal}")
        return index, conn

    print(f"[DEBUG][build_or_load_index] No existing index, building from scratch...")
    conn = init_db(db_path)
    index: faiss.Index | None = None
    next_id = 0

    line_iter = iter_lines(source_path)
    chunk_iter = iter_chunks_from_lines(line_iter, max_chars=MAX_CHARS_PER_CHUNK)

    for batch_num, (text_batch, emb_batch) in enumerate(
        embed_in_batches(chunk_iter, batch_size=EMBED_BATCH_SIZE), start=1
    ):
        if index is None:
            d = emb_batch.shape[1]
            print(f"[DEBUG][build_or_load_index] Initializing HNSW index | dim={d}")
            index = faiss.IndexHNSWFlat(d, 64)
            index.metric_type = faiss.METRIC_INNER_PRODUCT
            index.hnsw.efConstruction = 200
            index.hnsw.efSearch = 64

        index.add(emb_batch)
        insert_chunks(conn, next_id, text_batch)
        next_id += len(text_batch)
        print(f"[DEBUG][build_or_load_index] Batch {batch_num} added | total vectors so far={next_id}", flush=True)

    if index is None:
        raise RuntimeError("No chunks were generated; index is empty.")

    print(f"[DEBUG][build_or_load_index] Index build complete | total vectors={next_id}")
    print(f"[DEBUG][build_or_load_index] Saving index to '{index_path}'...")
    faiss.write_index(index, index_path)
    print(f"[DEBUG][build_or_load_index] Index saved to disk successfully")
    return index, conn


def retrieve_from_all(
    indexes: List[faiss.Index],
    conns: List[sqlite3.Connection],
    query: str,
    per_index_top_n: int = 4,
    final_top_n: int = 12,
) -> List[Tuple[str, float]]:
    print(f"[DEBUG][retrieve_from_all] Retrieving across {len(indexes)} index(es) | per_index_top_n={per_index_top_n} | final_top_n={final_top_n}")
    all_results: List[Tuple[str, float]] = []
    q_emb = embed_query(query)

    for i, (index, conn) in enumerate(zip(indexes, conns)):
        print(f"[DEBUG][retrieve_from_all] Searching index {i+1}/{len(indexes)} | vectors={index.ntotal}")
        sims, idxs = index.search(q_emb, per_index_top_n)
        for j, idx in enumerate(idxs[0]):
            if idx < 0:
                print(f"[DEBUG][retrieve_from_all]   slot {j}: invalid idx={idx}, skipping")
                continue
            text = fetch_chunk_by_id(conn, int(idx))
            score = float(sims[0][j])
            print(f"[DEBUG][retrieve_from_all]   slot {j}: idx={idx} | score={score:.4f} | preview='{text[:60]}'")
            all_results.append((text, score))

    all_results.sort(key=lambda x: x[1], reverse=True)
    trimmed = all_results[:final_top_n]
    print(f"[DEBUG][retrieve_from_all] Total results before trim={len(all_results)} | returning top {len(trimmed)}")
    for rank, (text, score) in enumerate(trimmed, 1):
        print(f"[DEBUG][retrieve_from_all]   rank {rank}: score={score:.4f} | '{text[:80]}'")
    return trimmed


def build_context_from_retrieval(retrieved: List[Tuple[str, float]]) -> str:
    print(f"[DEBUG][build_context_from_retrieval] Building context from {len(retrieved)} chunks")
    parts = []
    for i, (chunk, score) in enumerate(retrieved, start=1):
        parts.append(f"### {i}. (similarity: {score:.2f})\n{chunk}")
    context = "\n\n".join(parts)
    print(f"[DEBUG][build_context_from_retrieval] Context built | total chars={len(context)}")
    return context


@app.post("/query", response_model=QueryResponse)
async def query_files(
    files: List[UploadFile] = File(...),
    query: str = Form(...)
):
    print(f"\n{'='*60}")
    print(f"[DEBUG][/query] New request received")
    print(f"[DEBUG][/query] Query: '{query}'")
    print(f"[DEBUG][/query] Files: {[f.filename for f in files]}")
    print(f"{'='*60}")

    try:
        user_id = str(uuid.uuid4())
        print(f"[DEBUG][/query] Generated user_id: '{user_id}'")

        workspace = get_user_workspace(user_id)

        tabular_paths: List[str] = []
        rag_paths: List[str] = []

        for file in files:
            file_path = workspace / file.filename
            content = await file.read()
            with open(file_path, "wb") as f:
                f.write(content)
            print(f"[DEBUG][/query] Saved '{file.filename}' -> '{file_path}' ({len(content)} bytes)")

            path_str = str(file_path)
            if is_tabular(path_str):
                tabular_paths.append(path_str)
                print(f"[DEBUG][/query] '{file.filename}' routed to TABULAR pipeline")
            else:
                rag_paths.append(path_str)
                print(f"[DEBUG][/query] '{file.filename}' routed to RAG pipeline")

        print(f"[DEBUG][/query] Routing summary | tabular={len(tabular_paths)} file(s) | rag={len(rag_paths)} file(s)")
        answers: List[str] = []

        # ── TABULAR PIPELINE ──
        if tabular_paths:
            print(f"\n[DEBUG][/query] == TABULAR PIPELINE START ({len(tabular_paths)} file(s)) ==")
        for tab_path in tabular_paths:
            print(f"[DEBUG][/query] Processing tabular file: '{tab_path}'")
            df = load_dataframe(tab_path)
            schema_context = build_schema_context(df)
            answer = query_tabular_file(df, schema_context, query)
            print(f"[DEBUG][/query] Tabular answer for '{tab_path}': '{answer[:150]}'")
            answers.append(answer)
        if tabular_paths:
            print(f"[DEBUG][/query] == TABULAR PIPELINE END ==\n")

        # ── RAG PIPELINE ──
        if rag_paths:
            print(f"[DEBUG][/query] == RAG PIPELINE START ({len(rag_paths)} file(s)) ==")
            indexes, conns = build_user_indexes(user_id, rag_paths)

            print(f"[DEBUG][/query] Running retrieval across {len(indexes)} index(es)...")
            retrieved = retrieve_from_all(
                indexes=indexes,
                conns=conns,
                query=query,
                per_index_top_n=4,
                final_top_n=8,
            )

            context = build_context_from_retrieval(retrieved)

            instruction_prompt = (
                "You are a strict Retrieval-Augmented Generation (RAG) assistant.\n"
                "Your ONLY source of information is the context provided below.\n"
                "RULES:\n"
                "1. Do NOT use any outside knowledge or provide information not present in the context.\n"
                "2. If the answer is not explicitly stated in the context, you MUST say: "
                "'The provided documents do not contain enough information.'\n"
                "3. Do not mention that you are using context or documents; just answer the question.\n\n"
                f"Context:\n{context}\n\n"
            )

            print(f"[DEBUG][/query] Sending RAG prompt to LLM | model={LANGUAGE_MODEL} | context_chars={len(context)}")
            response = ollama.chat(
                model=LANGUAGE_MODEL,
                messages=[
                    {"role": "system", "content": instruction_prompt},
                    {"role": "user", "content": query},
                ],
            )
            rag_answer = response["message"]["content"]
            print(f"[DEBUG][/query] RAG answer preview: '{rag_answer[:150]}'")
            answers.append(rag_answer)

            print(f"[DEBUG][/query] Closing {len(conns)} DB connection(s)...")
            print(f"[DEBUG][/query] DB connections closed")
            print(f"[DEBUG][/query] == RAG PIPELINE END ==\n")

        final_answer = "\n\n---\n\n".join(answers) if answers else "No files were processed."
        print(f"[DEBUG][/query] Final answer length={len(final_answer)} chars")
        print(f"{'='*60}")
        print(f"[DEBUG][/query] Request complete")
        print(f"{'='*60}\n")
        return {"answer": final_answer}

    except Exception as e:
        print(f"[DEBUG][/query] EXCEPTION: {type(e).__name__}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


if __name__ == "__main__":
    import uvicorn
    print(f"[DEBUG][main] Starting FastAPI server on 0.0.0.0:8000")
    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=True)  